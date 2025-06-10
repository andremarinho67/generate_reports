import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import unittest
from generate_reports import tokenize_key_aspects, parse_docx, create_word
from docx import Document as DocxDocument

class TestGenerateReports(unittest.TestCase):
    def test_tokenize_key_aspects_empty(self):
        self.assertEqual(tokenize_key_aspects(""), [])

    def test_tokenize_key_aspects_single_line(self):
        self.assertEqual(tokenize_key_aspects("- Point 1 - Point 2"), ["Point 1", "Point 2"])

    def test_tokenize_key_aspects_multiline(self):
        s = "- Point 1\n- Point 2\n- Point 3"
        self.assertEqual(tokenize_key_aspects(s), ["Point 1", "Point 2", "Point 3"])

    def test_parse_docx(self):
        # Prepare a minimal DOCX file for testing
        from docx import Document
        doc = Document()
        doc.add_paragraph("Title: Test Entry\nDate: 2024-01-01\nCountry: Testland\nSummary: This is a summary.\nKey Aspects:\n- Aspect 1\n- Aspect 2\nLink: http://example.com\nAvailability: Public")
        test_docx = "test_input.docx"
        doc.save(test_docx)
        entries = parse_docx(test_docx)
        os.remove(test_docx)
        self.assertEqual(len(entries), 1)
        self.assertEqual(entries[0]["Title"], "Test Entry")
        self.assertIn("Aspect 1", entries[0]["Key Aspects"])

    def test_create_word_end_to_end(self):
        # Prepare entries
        entries = [{
            "Title": "Test Entry",
            "Date": "2024-01-01",
            "Country": "Testland",
            "Summary": "This is a summary.",
            "Key Aspects": "- Aspect 1\n- Aspect 2",
            "Link": "http://example.com",
            "Availability": "Public"
        }]
        output_docx = "test_output.docx"
        create_word(entries, output_docx)
        # Now check the output docx
        doc = DocxDocument(output_docx)
        # Gather all text from all tables
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + "\n"
        # Now assert
        self.assertIn("Test Entry", table_text)
        self.assertIn("2024-01-01", table_text)
        self.assertIn("Testland", table_text)
        self.assertIn("This is a summary.", table_text)
        self.assertIn("Aspect 1", table_text)
        self.assertIn("http://example.com", table_text)
        self.assertIn("Public", table_text)
        os.remove(output_docx)

    def test_table_structure(self):
        """Check that the generated table has the expected number of rows and columns."""
        entries = [{
            "Title": "Test Entry",
            "Date": "2024-01-01",
            "Country": "Testland",
            "Summary": "This is a summary.",
            "Key Aspects": "- Aspect 1\n- Aspect 2",
            "Link": "http://example.com",
            "Availability": "Public"
        }]
        output_docx = "test_output.docx"
        create_word(entries, output_docx)
        doc = DocxDocument(output_docx)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 4, "Table should have 4 rows")
        self.assertEqual(len(table.columns), 6, "Table should have 6 columns")
        os.remove(output_docx)

    def test_label_cells_are_bold(self):
        """Check that label cells have bold text."""
        entries = [{
            "Title": "Test Entry",
            "Date": "2024-01-01",
            "Country": "Testland",
            "Summary": "This is a summary.",
            "Key Aspects": "- Aspect 1\n- Aspect 2",
            "Link": "http://example.com",
            "Availability": "Public"
        }]
        output_docx = "test_output.docx"
        create_word(entries, output_docx)
        doc = DocxDocument(output_docx)
        table = doc.tables[0]
        label_cells = [(0, 0), (0, 2), (0, 4), (1, 0), (2, 0), (3, 0)]
        for row, col in label_cells:
            cell = table.cell(row, col)
            self.assertTrue(
                any(run.bold for run in cell.paragraphs[0].runs),
                f"Label cell ({row},{col}) is not bold"
            )
        os.remove(output_docx)

    def test_summary_cell_no_leading_empty_line(self):
        """Check that the Summary content does not start with an empty line."""
        entries = [{
            "Title": "Test Entry",
            "Date": "2024-01-01",
            "Country": "Testland",
            "Summary": "This is a summary.",
            "Key Aspects": "- Aspect 1\n- Aspect 2",
            "Link": "http://example.com",
            "Availability": "Public"
        }]
        output_docx = "test_output.docx"
        create_word(entries, output_docx)
        doc = DocxDocument(output_docx)
        table = doc.tables[0]
        summary_cell = table.cell(1, 1)
        if summary_cell.paragraphs:
            first_para = summary_cell.paragraphs[0].text
            self.assertTrue(first_para.strip(), "Summary content starts with an empty line or is empty")
            self.assertFalse(first_para.startswith('\n'), "Summary content starts with a newline")
            self.assertFalse(first_para.startswith('\r'), "Summary content starts with a carriage return")
        os.remove(output_docx)

    def test_merged_cells_for_summary_link_availability(self):
        """Check that summary, link, and availability value cells are merged across columns 1-5."""
        entries = [{
            "Title": "Test Entry",
            "Date": "2024-01-01",
            "Country": "Testland",
            "Summary": "This is a summary.",
            "Key Aspects": "- Aspect 1\n- Aspect 2",
            "Link": "http://example.com",
            "Availability": "Public"
        }]
        output_docx = "test_output.docx"
        create_word(entries, output_docx)
        doc = DocxDocument(output_docx)
        table = doc.tables[0]
        for row in [1, 2, 3]:
            merged_texts = [table.cell(row, col).text for col in range(1, 6)]
            self.assertTrue(
                all(text == merged_texts[0] for text in merged_texts),
                f"Row {row} value cells are not properly merged"
            )
        os.remove(output_docx)

if __name__ == "__main__":
    unittest.main()