import unittest
from docx import Document as DocxDocument

class TestWordTableContent(unittest.TestCase):
    def setUp(self):
        # Load the generated Word document
        self.doc = DocxDocument("output.docx")
        self.table = self.doc.tables[0]

    def test_table_cells_content(self):
        """
        Check that the main table cells contain the expected text.
        Adjust the expected values as needed for your actual output.docx.
        """
        expected = {
            (0, 0): "Title",
            (0, 2): "Date",
            (0, 4): "Country",
            (1, 0): "Summary",
            (2, 0): "Link",
            (3, 0): "Availability",
        }
        # Check label cells
        for (row, col), exp_text in expected.items():
            actual = self.table.cell(row, col).text.strip()
            self.assertEqual(actual, exp_text, f"Cell ({row},{col}) expected '{exp_text}' but got '{actual}'")

        # Check that value cells are not empty (except country value cell if flag is used)
        self.assertTrue(self.table.cell(0, 1).text.strip(), "Title value cell is empty")
        self.assertTrue(self.table.cell(0, 3).text.strip(), "Date value cell is empty")
        # Country value cell may be empty if a flag image is used, so skip or warn
        # self.assertTrue(self.table.cell(0, 5).text.strip(), "Country value cell is empty")
        self.assertTrue(self.table.cell(1, 1).text.strip(), "Summary value cell is empty")
        self.assertTrue(self.table.cell(2, 1).text.strip(), "Link value cell is empty")
        self.assertTrue(self.table.cell(3, 1).text.strip(), "Availability value cell is empty")

    def test_label_cells_are_bold(self):
        """Check that label cells have bold text."""
        label_cells = [(0, 0), (0, 2), (0, 4), (1, 0), (2, 0), (3, 0)]
        for row, col in label_cells:
            cell = self.table.cell(row, col)
            # Check if any run in the first paragraph is bold
            self.assertTrue(
                any(run.bold for run in cell.paragraphs[0].runs),
                f"Label cell ({row},{col}) is not bold"
            )

    def test_summary_cell_contains_key_aspects(self):
        """Check that the summary cell contains expected key aspects."""
        summary_text = self.table.cell(1, 1).text
        self.assertIn("Key", summary_text)  # Adjust as needed for your data

    def test_merged_cells_for_summary_link_availability(self):
        """Check that summary, link, and availability value cells are merged across columns 1-5."""
        for row in [1, 2, 3]:
            merged_texts = [self.table.cell(row, col).text for col in range(1, 6)]
            self.assertTrue(
                all(text == merged_texts[0] for text in merged_texts),
                f"Row {row} value cells are not properly merged"
            )

    def test_flag_cell_contains_image_or_text(self):
        """Check that the country value cell contains either text or an image (drawing)."""
        cell = self.table.cell(0, 5)
        has_text = bool(cell.text.strip())
        # Check for image by looking for <w:drawing> in the cell XML
        has_image = "w:drawing" in cell._tc.xml
        self.assertTrue(
            has_text or has_image,
            "Country value cell contains neither text nor image"
        )

    def test_all_value_cells_nonempty(self):
        """Check that all value cells except possibly the flag cell are non-empty."""
        value_cells = [(0, 1), (0, 3), (1, 1), (2, 1), (3, 1)]
        for row, col in value_cells:
            self.assertTrue(
                self.table.cell(row, col).text.strip(),
                f"Value cell ({row},{col}) is empty"
            )

    def test_summary_content_no_leading_empty_line(self):
        """Check that the Summary content does not start with an empty line."""
        summary_cell = self.table.cell(1, 1)
        # Get the first paragraph text (should be the summary)
        if summary_cell.paragraphs:
            first_para = summary_cell.paragraphs[0].text
            # It should not be empty or start with a newline
            self.assertTrue(first_para.strip(), "Summary content starts with an empty line or is empty")
            self.assertFalse(first_para.startswith('\n'), "Summary content starts with a newline")
            self.assertFalse(first_para.startswith('\r'), "Summary content starts with a carriage return")

    def test_table_count_matches_entries(self):
        """Ensure the number of tables matches the number of entries in the report."""
        # Adjust this number to match your expected number of entries
        expected_table_count = 9  # Or dynamically determine from your input
        self.assertEqual(len(self.doc.tables), expected_table_count)

if __name__ == "__main__":
    unittest.main()