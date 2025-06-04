from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Image, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import ListFlowable, ListItem, KeepTogether, Flowable
import os
import re
import argparse

# Custom One Consulting blue color
one_consult_blue = colors.Color(87 / 255, 155 / 255, 156 / 255, 1)  # RGBA
transparent_blue = colors.Color(87 / 255, 155 / 255, 156 / 255, 0.15)  # RGBA

# Map country to flag image filenames (ensure these files exist)
country_flags = {
    'Luxembourg': 'flags/luxembourg.png',
    'Ireland': 'flags/ireland.png',
    'UK': 'flags/uk.png',
    'Switzerland': 'flags/switzerland.png',
    'European Union': 'flags/european_union.png'
}

styles = getSampleStyleSheet()
normal_style = styles['Normal']
normal_style.fontName = 'Helvetica'
normal_style.fontSize = 10
normal_style.leading = 12

label_style = ParagraphStyle(
    'LabelStyle',
    parent=normal_style,
    backColor=one_consult_blue,
    fontName='Helvetica-Bold',
    fontSize=10,
    alignment=1,  # center align
    spaceAfter=4,
)

value_style = ParagraphStyle(
    'ValueStyle',
    parent=normal_style,
    fontName='Helvetica',
    fontSize=10,
    leading=12,
)


def parse_docx(file_path):
    print(f"[parse_docx] Loading document: {file_path}")
    document = DocxDocument(file_path)
    full_text = ' '.join(
        [p.text.strip() for p in document.paragraphs if p.text.strip()])
    print(f"[parse_docx] Full text length: {len(full_text)} chars")

    raw_entries = re.split(r'(?=Title:)', full_text)
    print(
        f"[parse_docx] Found {len(raw_entries)} raw entries (including empty first)"
    )

    entries = []
    # Regex pattern: Key Aspects is a single string (may be empty)
    pattern = (
        r'Title:\s*(.*?)\s+Date:\s*(.*?)\s+Country:\s*(.*?)\s+Summary:\s*(.*?)'
        r'(?:\s+Key Aspects:\s*((?:- .*\s*)*))?'
        r'Link:\s*(.*?)\s+Availability:\s*(.*)')

    for idx, entry_text in enumerate(raw_entries):
        entry_text = entry_text.strip()
        if not entry_text:
            continue
        print(f"[parse_docx] Parsing entry #{idx+1}")

        match = re.match(pattern, entry_text, re.DOTALL)
        if not match:
            print(
                f"  [WARNING] Entry #{idx+1} does not match expected format!")
            print(f"  Text: {entry_text[:100]}...")
            continue

        key_aspects_raw = match.group(5)
        if key_aspects_raw:
            key_aspects = key_aspects_raw.strip()
        else:
            key_aspects = ""

        entry = {
            "Title": match.group(1).strip(),
            "Date": match.group(2).strip(),
            "Country": match.group(3).strip(),
            "Summary": match.group(4).strip(),
            "Key Aspects": key_aspects,
            "Link": match.group(6).strip(),
            "Availability": match.group(7).strip(),
        }

        print(
            f"  Parsed Key Aspects for '{entry['Title']}': {entry['Key Aspects'][:50]}{'...' if len(entry['Key Aspects'])>50 else ''}"
        )

        entries.append(entry)

    print(f"[parse_docx] Completed parsing entries. Total: {len(entries)}")
    return entries


def tokenize_key_aspects(key_aspects_str):
    """Split Key Aspects string into a list of bullet points, robust to both multi-line and single-line formats."""
    if not key_aspects_str:
        return []

    # Normalize into a single string (in case it's multi-line)
    combined = " ".join(line.strip() for line in key_aspects_str.splitlines())

    # Split on ' - ' and clean each bullet
    parts = [
        part.strip(" -") for part in combined.split(" - ") if part.strip(" -")
    ]
    return parts


def build_table_for_entry(entry):
    print(
        f"[build_table_for_entry] Building table for entry: {entry['Title']}")

    # Prepare Title split for two lines max
    title = entry["Title"]
    if len(title) > 40:
        parts = title.split(' ')
        mid = len(parts) // 2
        title_text = ' '.join(parts[:mid]) + '\n' + ' '.join(parts[mid:])
    else:
        title_text = title

    # Prepare flag image if exists
    flag_path = country_flags.get(entry["Country"])
    flag_img = None
    if flag_path and os.path.isfile(flag_path):
        try:
            if entry["Country"] == "Switzerland":
                # Square flag for Switzerland
                flag_img = Image(flag_path,
                                 width=0.5 * inch,
                                 height=0.5 * inch)
            else:
                # Wider flags for others, keep aspect ratio approx 5:3
                flag_img = Image(flag_path,
                                 width=0.5 * inch,
                                 height=0.3 * inch)
        except Exception as e:
            print(f"  Error loading flag image: {e}")
            flag_img = None
    else:
        print(
            f"  No flag image found for country '{entry['Country']}' at '{flag_path}'"
        )

    # Wrap flag image in a Table cell for centering
    if flag_img:
        flag_img = Table([[flag_img]],
                         colWidths=[0.5 * inch],
                         rowHeights=[0.5 * inch])
        flag_img.setStyle(
            TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))

    # Increase font leading for multiline wrapping text to avoid overlap
    summary_style = ParagraphStyle(
        'SummaryStyle',
        parent=value_style,
        leading=16,
    )

    # Compose summary with key aspects
    summary_text = entry.get('Summary', '') or ''
    summary_flowables = [Paragraph(summary_text, summary_style)]

    key_aspects_list = tokenize_key_aspects(entry.get("Key Aspects", ""))
    bullet_items = [
        ListItem(Paragraph(point, value_style)) for point in key_aspects_list
        if point
    ]

    if bullet_items:
        summary_flowables.append(Spacer(1, 6))
        summary_flowables.append(Paragraph('<b>Key Aspects:</b>', value_style))
        summary_flowables.append(
            ListFlowable(bullet_items, bulletType='bullet', leftIndent=12))

    # Remove any None values
    summary_flowables = [f for f in summary_flowables if f is not None]
    summary_cell_content = summary_flowables if len(
        summary_flowables) > 1 else summary_flowables[0]

    # Table data (6 columns)
    data = [
        [
            Paragraph('<b>Title</b>', label_style),
            Paragraph(title_text, value_style),
            Paragraph('<b>Date</b>', label_style),
            Paragraph(entry['Date'], value_style),
            Paragraph('<b>Country</b>', label_style),
            flag_img if flag_img else Paragraph(entry['Country'], value_style)
        ],
        [
            Paragraph('<b>Summary</b>', label_style), summary_cell_content, '',
            '', '', ''
        ],
        [
            Paragraph('<b>Link</b>', label_style),
            Paragraph(entry['Link'], value_style), '', '', '', ''
        ],
        [
            Paragraph('<b>Availability</b>', label_style),
            Paragraph(entry['Availability'], value_style), '', '', '', ''
        ],
    ]

    # Set row heights: fixed except summary auto
    base_height = 14
    row_heights = [
        base_height * 3,  # Title row
        None,  # Summary auto height
        base_height * 3,  # Link row
        base_height * 3  # Availability row
    ]

    # --- Modern semi-transparent border color ---
    semi_transparent_border = colors.Color(87 / 255, 155 / 255, 156 / 255,
                                           0.4)  # RGBA

    t = Table(data,
              colWidths=[
                  1 * inch, 2.5 * inch, 0.8 * inch, 1 * inch, 0.8 * inch,
                  1 * inch
              ],
              rowHeights=row_heights)
    style = TableStyle([
        # Modern semi-transparent borders
        ('BOX', (0, 0), (-1, -1), 1, semi_transparent_border),
        ('INNERGRID', (0, 0), (-1, -1), 1, semi_transparent_border),
        # Label backgrounds (solid)
        ('BACKGROUND', (0, 0), (0, 0), one_consult_blue),
        ('BACKGROUND', (2, 0), (2, 0), one_consult_blue),
        ('BACKGROUND', (4, 0), (4, 0), one_consult_blue),
        ('BACKGROUND', (0, 1), (0, 1), one_consult_blue),
        ('BACKGROUND', (0, 2), (0, 2), one_consult_blue),
        ('BACKGROUND', (0, 3), (0, 3), one_consult_blue),
        # Content backgrounds (transparent)
        ('BACKGROUND', (1, 0), (1, 0), transparent_blue),
        ('BACKGROUND', (3, 0), (3, 0), transparent_blue),
        ('BACKGROUND', (5, 0), (5, 0), transparent_blue),
        ('BACKGROUND', (1, 1), (-1, 1), transparent_blue),
        ('BACKGROUND', (1, 2), (-1, 2), transparent_blue),
        ('BACKGROUND', (1, 3), (-1, 3), transparent_blue),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (1, 0), (1, 0), 'LEFT'),
        ('ALIGN', (1, 1), (1, 3), 'LEFT'),
        ('ALIGN', (5, 0), (5, 0), 'CENTER'),
        ('SPAN', (1, 1), (-1, 1)),
        ('SPAN', (1, 2), (-1, 2)),
        ('SPAN', (1, 3), (-1, 3)),
    ])
    t.setStyle(style)

    print(f"  Table built for entry: {entry['Title']}")
    return t


def set_cell_background(cell, color_hex):
    if color_hex is None:
        return
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
        nsdecls('w'), color_hex))
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove existing shading if any
    for child in tc_pr.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd'
    ):
        tc_pr.remove(child)
    tc_pr.append(shading_elm)


def create_word(entries, output_docx):
    print("[create_word] Creating Word document:", output_docx)
    doc = DocxDocument()

    for i, entry in enumerate(entries, 1):
        print(f"[create_word] Processing entry #{i}")

        table = doc.add_table(rows=4, cols=6)
        table.autofit = False
        widths = [
            Inches(1),
            Inches(2.5),
            Inches(0.8),
            Inches(1),
            Inches(0.8),
            Inches(1)
        ]
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = width

        # Prepare Title split
        title = entry["Title"]
        if len(title) > 40:
            parts = title.split(' ')
            mid = len(parts) // 2
            title_text = ' '.join(parts[:mid]) + '\n' + ' '.join(parts[mid:])
        else:
            title_text = title

        # Row 1
        cells = [
            (table.cell(0, 0), "Title"),
            (table.cell(0, 1), title_text),
            (table.cell(0, 2), "Date"),
            (table.cell(0, 3), entry["Date"]),
            (table.cell(0, 4), "Country"),
            (table.cell(0, 5), None),  # Flag or country text filled later
        ]

        # Set backgrounds for label/content fields in Row 1
        for idx, (cell, text) in enumerate(cells):
            if text in ["Title", "Date", "Country"]:
                set_cell_background(cell, "579B9C")  # one_consult_blue
            else:
                set_cell_background(cell, "E2F3F3")  # transparent_blue
            if text is not None:
                cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Make labels bold (in label cells)
        for idx in [0, 2, 4]:
            cell = table.cell(0, idx)
            cell.paragraphs[0].runs[0].font.bold = True

        # Add flag or country text centered in last cell
        country_cell_value = table.cell(0, 5)
        flag_path = country_flags.get(entry["Country"])
        if flag_path and os.path.isfile(flag_path):
            # Remove all paragraphs except the first
            while len(country_cell_value.paragraphs) > 1:
                p = country_cell_value.paragraphs[-1]
                p._element.getparent().remove(p._element)
            # Clear any existing text
            country_cell_value.text = ""
            paragraph = country_cell_value.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_picture(flag_path, width=Inches(0.5))
            # Set vertical alignment after adding the image
            country_cell_value.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            country_cell_value.text = entry["Country"]
            country_cell_value.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Row 2 Summary
        summary_label_cell = table.cell(1, 0)
        summary_label_cell.text = "Summary"
        set_cell_background(summary_label_cell, "579B9C")  # one_consult_blue
        summary_label_cell.paragraphs[0].runs[0].font.bold = True
        summary_label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        summary_value_cell = table.cell(1, 1)
        summary_value_cell.text = ""
        set_cell_background(summary_value_cell, "E2F3F3")  # transparent_blue
        summary_text = entry["Summary"].lstrip('\n').lstrip()
        p = summary_value_cell.add_paragraph(summary_text)
        p.paragraph_format.space_after = Pt(6)
        key_aspects_list = tokenize_key_aspects(entry.get("Key Aspects", ""))
        if key_aspects_list:
            p = summary_value_cell.add_paragraph()
            run = p.add_run("Key Aspects:")
            run.bold = True
            for point in key_aspects_list:
                bullet = summary_value_cell.add_paragraph(point,
                                                          style='List Bullet')
                bullet.paragraph_format.left_indent = Pt(18)
        summary_value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for col in range(2, 6):
            cell = table.cell(1, col)
            set_cell_background(cell, "E2F3F3")  # transparent_blue
            summary_value_cell.merge(cell)

        # Row 3 Link
        link_label_cell = table.cell(2, 0)
        link_label_cell.text = "Link"
        set_cell_background(link_label_cell, "579B9C")  # one_consult_blue
        link_label_cell.paragraphs[0].runs[0].font.bold = True
        link_label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        link_value_cell = table.cell(2, 1)
        link_value_cell.text = entry["Link"]
        set_cell_background(link_value_cell, "E2F3F3")  # transparent_blue
        link_value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for col in range(2, 6):
            cell = table.cell(2, col)
            set_cell_background(cell, "E2F3F3")  # transparent_blue
            link_value_cell.merge(cell)

        # Row 4 Availability
        avail_label_cell = table.cell(3, 0)
        avail_label_cell.text = "Availability"
        set_cell_background(avail_label_cell, "579B9C")  # one_consult_blue
        avail_label_cell.paragraphs[0].runs[0].font.bold = True
        avail_label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        avail_value_cell = table.cell(3, 1)
        avail_value_cell.text = entry["Availability"]
        set_cell_background(avail_value_cell, "E2F3F3")  # transparent_blue
        avail_value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for col in range(2, 6):
            cell = table.cell(3, col)
            set_cell_background(cell, "E2F3F3")  # transparent_blue
            avail_value_cell.merge(cell)

        # Set fixed height for the first row
        table.rows[0].height = Inches(0.6)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # --- Add modern light blue borders to all cells ---
        # Word does not support transparency, so use a light blue color
        border_color = "B3E0E2"  # Light blue hex (matches the blue, but lighter)
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Set all borders (top, left, bottom, right)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border_tag = f'w:{border_name}'
                    border = tcPr.find(
                        f'.//{border_tag}',
                        namespaces={
                            'w':
                            'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        })
                    if border is None:
                        from docx.oxml import OxmlElement
                        border = OxmlElement(border_tag)
                        tcPr.append(border)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '6')  # Thin border
                    border.set(qn('w:color'), border_color)
                    border.set(qn('w:space'), '0')

        # Add page break after each table except last
        if i != len(entries):
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

    doc.save(output_docx)
    print("[create_word] Document saved successfully.")


def create_pdf(entries, output_pdf):
    print("[create_pdf] Creating PDF document:", output_pdf)
    doc = SimpleDocTemplate(output_pdf,
                            pagesize=A4,
                            rightMargin=36,
                            leftMargin=36,
                            topMargin=36,
                            bottomMargin=36)

    elements = []

    for i, entry in enumerate(entries, 1):
        elements.append(build_table_for_entry(entry))
        if i != len(entries):
            elements.append(Spacer(1, 0.2 * inch))
            elements.append(PageBreak())

    doc.build(elements)
    print("[create_pdf] PDF saved successfully.")


def main():
    parser = argparse.ArgumentParser(
        description='Generate report from Word to PDF or Word.')
    parser.add_argument('input_docx', help='Input DOCX file with entries')
    parser.add_argument('-o',
                        '--output',
                        default='output.pdf',
                        help='Output filename (default: output.pdf)')
    parser.add_argument('-w',
                        '--word',
                        action='store_true',
                        help='Create a Word document instead of PDF')
    args = parser.parse_args()

    entries = parse_docx(args.input_docx)

    if args.word:
        # Force output to .docx extension if not provided
        if not args.output.lower().endswith('.docx'):
            output_docx = os.path.splitext(args.output)[0] + '.docx'
        else:
            output_docx = args.output
        create_word(entries, output_docx)
    else:
        # Force output to .pdf extension if not provided
        if not args.output.lower().endswith('.pdf'):
            output_pdf = os.path.splitext(args.output)[0] + '.pdf'
        else:
            output_pdf = args.output
        create_pdf(entries, output_pdf)


if __name__ == "__main__":
    main()
